import streamlit as st
try:
    import google.generativeai as genai
except ImportError:
    st.error("Missing 'google-generativeai'. Install it using: `pip install google-generativeai`")
    st.stop()
try:
    from PyPDF2 import PdfReader
except ImportError:
    st.error("Missing 'PyPDF2'. Install it using: `pip install PyPDF2`")
    st.stop()
try:
    import pytesseract
except ImportError:
    st.error("Missing 'pytesseract'. Install it using: `pip install pytesseract`. Also ensure Tesseract OCR is installed and added to PATH.")
    st.stop()
try:
    from PIL import Image
except ImportError:
    st.error("Missing 'Pillow'. Install it using: `pip install Pillow`")
    st.stop()
try:
    import numpy as np
except ImportError:
    st.error("Missing 'numpy'. Install it using: `pip install numpy`")
    st.stop()
try:
    import faiss
except ImportError:
    st.error("Missing 'faiss'. Install it using: `pip install faiss-cpu`. Ensure Microsoft Visual C++ Build Tools are installed.")
    st.stop()
try:
    import docx
except ImportError:
    st.error("Missing 'python-docx'. Install it using: `pip install python-docx`")
    st.stop()
try:
    import pandas as pd
except ImportError:
    st.error("Missing 'pandas'. Install it using: `pip install pandas`")
    st.stop()
try:
    from dotenv import load_dotenv
except ImportError:
    st.error("Missing 'python-dotenv'. Install it using: `pip install python-dotenv`")
    st.stop()
try:
    import pdfplumber
except ImportError:
    st.error("Missing 'pdfplumber'. Install it using: `pip install pdfplumber`")
    st.stop()
try:
    from transformers import CLIPProcessor, CLIPModel
except ImportError:
    st.error("Missing 'transformers'. Install it using: `pip install transformers torch`")
    st.stop()
import torch
import base64
import os
from datetime import datetime
import io

# Load environment variables
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    st.error("GEMINI_API_KEY not found in .env file. Please create a .env file in `C:\\Users\\bharg\\OneDrive\\Desktop\\ces_main` with: `GEMINI_API_KEY=your-api-key`")
    st.stop()

genai.configure(api_key=API_KEY)

# Initialize CLIP model
clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")

# Set Streamlit layout to wide for desktop, but allow responsive CSS for mobile
st.set_page_config(layout="wide")

# Sidebar Instructions
st.sidebar.title("Setup Instructions")
st.sidebar.markdown("""
**Required Setup**:
1. Save the following as `requirements.txt` in your project folder and run `pip install -r requirements.txt`:
2. Create a `.env` file in `C:\\Users\\bharg\\OneDrive\\Desktop\\ces_main` with:
3. Install Tesseract OCR: [Download](https://github.com/UB-Mannheim/tesseract/wiki) and add to PATH (e.g., `C:\\Program Files\\Tesseract-OCR`).
4. Install Microsoft Visual C++ Build Tools: [Download](https://visualstudio.microsoft.com/visual-cpp-build-tools/) and select 'Desktop development with C++'.
5. Place your logo at `C:\\Users\\bharg\\OneDrive\\Desktop\\ces_main\\logo.png`.
6. Run the app: `python -m streamlit run app.py`
""")

QNA_PROMPT = """
You are an advanced Retrieval-Augmented Generation (RAG) assistant for enterprise document analysis. 
Your task is to deliver only the most relevant information in a highly compressed, easy-to-read format. 
Strictly follow the response structure and rules below to ensure clear, crisp outputs.

Response Rules:
1. **Answer** – Begin with a brief restatement of the user’s question for context, then provide the answer in 1–2 short, direct sentences (≤2 lines). Compress all retrieved content, including text from documents or image descriptions (e.g., [Image on page X]), into minimal words while keeping essential details. Cite document sections or image metadata if available.  
2. **Recommendations and Insights** – Provide exactly 1–2 bullet points. Each must be ≤2 lines, compressing all insights into concise, executive-ready statements. Clearly distinguish between:
   - Facts taken directly from documents or image content.
   - Inferences, best practices, or external knowledge (explicitly labeled as such).

Guidelines:
- Prioritize brevity and clarity: every response must be instantly scannable.  
- Avoid filler, repetition, or verbose explanations.  
- Use simple, professional language for readability.  
- Synthesize overlapping content from documents and images into a single compact statement.  
- If information is unavailable, state clearly: "The document or image does not provide this information."  
- Do not deviate from the specified format under any circumstance.
"""

REPORT_PROMPT = """
You are an elite consulting assistant producing Big 4–style executive reports (KPMG, Deloitte, PwC, EY). 
Transform the provided documents and images into a professional, fact-based, and insight-driven analysis that is concise, structured, and decision-focused.

Structure:
1. **Executive Summary** – High-level overview in minimal lines, capturing critical takeaways from documents and images (e.g., [Image on page X]).
2. **Key Insights and Findings** – Condensed, data-driven highlights in bullets or tables, including image-derived insights.
3. **Risk Assessment & Opportunities** – Summarize risks, compliance issues, and opportunities in 1–3 lines, including image-based insights.
4. **Strategic Recommendations** – 2–3 actionable recommendations, each ≤2 lines, marking fact-based (from documents/images) vs. inferences.
5. **Supporting Evidence** – Cite document excerpts or image descriptions (e.g., [Image on page X]) in 1 line per point.

Guidelines:
- Compress multi-line content into minimal statements while preserving key meaning.  
- No filler or repetition; every line must deliver decision-making value.  
- Maintain a professional, neutral, authoritative tone.  
- Use bullets/tables for clarity.  
- State if information is missing and suggest next steps.  
- Output must be executive-ready: short, fact-based, and insightful.
"""

def extract_text_from_file(uploaded_file):
    file_type = uploaded_file.type
    try:
        if "pdf" in file_type:
            items = []  # List of {'type': 'text' or 'image', 'content': str or bytes, 'metadata': dict}
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    if page_text:
                        items.append({
                            'type': 'text',
                            'content': page_text,
                            'metadata': {'page': page_num}
                        })
                    # Extract images with positions
                    for img_idx, img in enumerate(page.images):
                        try:
                            x0, top, x1, bottom = img['x0'], img['top'], img['x1'], img['bottom']
                            # Validate bounding box
                            if x0 >= x1 or top >= bottom or x0 < 0 or top < 0 or x1 > page.width or bottom > page.height:
                                st.warning(f"Skipping image {img_idx} on page {page_num}: Invalid bounding box ({x0},{top},{x1},{bottom})")
                                continue
                            cropped_page = page.crop((x0, top, x1, bottom))
                            pil_img = cropped_page.to_image().original
                            buffer = io.BytesIO()
                            pil_img.save(buffer, format="PNG")
                            image_bytes = buffer.getvalue()
                            position = (x0, top, x1, bottom)
                            items.append({
                                'type': 'image',
                                'content': image_bytes,
                                'metadata': {'page': page_num, 'position': position, 'img_idx': img_idx}
                            })
                        except Exception as e:
                            st.warning(f"Error extracting image {img_idx} on page {page_num}: {str(e)}")
            return items
        elif "image" in file_type:
            image = Image.open(uploaded_file)
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            return [{'type': 'image', 'content': buffer.getvalue(), 'metadata': {}}]
        elif "text" in file_type:
            text = uploaded_file.read().decode("utf-8")
            return [{'type': 'text', 'content': text, 'metadata': {}}]
        elif "csv" in file_type:
            df = pd.read_csv(uploaded_file)
            return [{'type': 'text', 'content': df.to_string(index=False), 'metadata': {}}]
        elif ("excel" in file_type) or uploaded_file.name.endswith((".xls", ".xlsx")):
            df = pd.read_excel(uploaded_file)
            return [{'type': 'text', 'content': df.to_string(index=False), 'metadata': {}}]
        else:
            return "Unsupported file type."
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return ""

def chunk_text(text, chunk_size=500):
    if not text or text == "Unsupported file type.":
        return []
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunks.append(" ".join(words[i:i + chunk_size]))
    return chunks

def get_embeddings(items):
    embeddings = []
    for item in items:
        try:
            if item['type'] == 'text':
                inputs = clip_processor(text=[item['content']], return_tensors="pt", padding=True, truncation=True)
                with torch.no_grad():
                    emb = clip_model.get_text_features(**inputs).cpu().numpy()
                    if emb.shape[0] != 1:
                        st.warning(f"Skipping text item: Unexpected embedding shape {emb.shape}")
                        continue
                    emb = emb[0]
            elif item['type'] == 'image':
                pil_img = Image.open(io.BytesIO(item['content']))
                inputs = clip_processor(images=[pil_img], return_tensors="pt")  # Explicit single image
                with torch.no_grad():
                    emb = clip_model.get_image_features(**inputs).cpu().numpy()
                    if emb.shape[0] != 1:
                        st.warning(f"Skipping image item: Unexpected embedding shape {emb.shape}")
                        continue
                    emb = emb[0]
            embeddings.append(emb)
        except Exception as e:
            st.warning(f"Error generating embedding for item: {str(e)}")
            continue
    if not embeddings:
        st.error("No valid embeddings generated.")
        return None
    return np.array(embeddings, dtype=np.float32)

def create_faiss_index(embeddings):
    try:
        dimension = embeddings.shape[1]  # CLIP: 512
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)
        return index
    except Exception as e:
        st.error(f"Error creating FAISS index: {str(e)}")
        return None

def retrieve_relevant_chunks(query, index, items, chunk_metadata, k=5):
    try:
        inputs = clip_processor(text=[query], return_tensors="pt", padding=True, truncation=True)
        with torch.no_grad():
            query_embedding = clip_model.get_text_features(**inputs).cpu().numpy()
        distances, indices = index.search(query_embedding, k)
        relevant = []
        for idx in indices[0]:
            if idx < len(items):
                relevant.append((items[idx], chunk_metadata[idx]))
        return relevant
    except Exception as e:
        st.error(f"Error retrieving relevant chunks: {str(e)}")
        return []

def save_report_to_word(report_text, filename_prefix="Generated_Analysis_Report"):
    try:
        doc = docx.Document()
        doc.add_heading("Generated Analysis Report", 0)
        for line in report_text.split("\n"):
            if line.startswith("1) ") or line.startswith("2) ") or line.startswith("3) ") or line.startswith("4) ") or line.startswith("5) "):
                doc.add_heading(line, level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line.replace("- ", ""), style="ListBullet")
            else:
                doc.add_paragraph(line)
        filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error saving docx: {str(e)}")
        return None

# Streamlit App Layout
LOGO_PATH = "logo.png"
with st.container():
    col1, col2 = st.columns([1, 6])
    with col1:
        if os.path.exists(LOGO_PATH):
            st.markdown(
                """
                <style>
                @media (max-width: 600px) {
                    .logo-img {max-width: 48px !important; height: auto !important;}
                }
                @media (min-width: 601px) {
                    .logo-img {max-width: 72px !important; height: auto !important;}
                }
                </style>
                """,
                unsafe_allow_html=True
            )
            st.image(LOGO_PATH, output_format="auto", caption=None, use_container_width=False, clamp=False, channels="RGB", width=72)
        else:
            st.warning("Logo file not found at 'logo.png'. Please place your logo in `C:\\Users\\bharg\\OneDrive\\Desktop\\ces_main`.")
    with col2:
        st.markdown(
            """
            <style>
            .responsive-heading {
                font-size: clamp(1.1rem, 4vw, 2.2rem);
                text-align: left;
                margin-top: 0;
                margin-bottom: 0;
                white-space: nowrap;
                overflow: hidden;
                text-overflow: ellipsis;
                font-weight: 700;
                letter-spacing: -1px;
            }
            @media (max-width: 600px) {
                .responsive-heading {
                    font-size: 1.1rem !important;
                }
            }
            </style>
            <h1 class="responsive-heading">PDF Analyst Chatbot and Report Generator</h1>
            """,
            unsafe_allow_html=True
        )

st.markdown(
    """
    <style>
    .responsive-subheader {
        font-size: clamp(0.9rem, 2.5vw, 1.3rem);
        margin-top: 10px;
        margin-bottom: 0;
    }
    </style>
    <div class="responsive-subheader">Upload files (PDF, CSV, XLSX) for Q&amp;A and Generate Analysis Reports</div>
    """,
    unsafe_allow_html=True
)

# File Uploader
uploaded_files = st.file_uploader("Upload files (PDF, CSV, XLSX)", type=["pdf", "csv", "xls", "xlsx"], accept_multiple_files=True)

# Store extracted document contents and chunks
items_list = []
chunk_metadata = []
faiss_index = None
if uploaded_files:
    for file_idx, uploaded_file in enumerate(uploaded_files):
        items = extract_text_from_file(uploaded_file)
        if items and items != "Unsupported file type.":
            for item_idx, item in enumerate(items):
                item['metadata']['file'] = uploaded_file.name
                item['metadata']['global_idx'] = len(items_list)
                items_list.append(item)
                chunk_metadata.append((uploaded_file.name, item_idx))
    if items_list:
        final_items = []
        final_metadata = []
        item_idx = 0
        for item in items_list:
            if item['type'] == 'text':
                text_chunks = chunk_text(item['content'])
                for chunk_idx, chunk in enumerate(text_chunks):
                    final_items.append({
                        'type': 'text',
                        'content': chunk,
                        'metadata': {
                            'file': item['metadata']['file'],
                            'page': item['metadata'].get('page', 0),
                            'chunk_idx': chunk_idx
                        }
                    })
                    final_metadata.append((item['metadata']['file'], item_idx))
                    item_idx += 1
            else:  # Image
                final_items.append(item)
                final_metadata.append((item['metadata']['file'], item_idx))
                item_idx += 1
        items_list = final_items
        chunk_metadata = final_metadata

        embeddings = get_embeddings(items_list)
        if embeddings is not None and embeddings.size > 0:
            faiss_index = create_faiss_index(embeddings)
            if faiss_index:
                st.success("Files indexed successfully in Vector DB (multimodal).")
            else:
                st.error("Failed to create Vector DB index. Check error messages above.")
                st.stop()
        else:
            st.error("No valid embeddings generated. Check warnings for skipped items.")
            st.stop()

# Generate Analysis Report Button
if st.button("Generate Report") and uploaded_files:
    with st.spinner("Generating Report..."):
        report_query = "Key insights, risks, and recommendations from the documents"
        relevant = retrieve_relevant_chunks(report_query, faiss_index, items_list, chunk_metadata, k=10)
        content_list = [REPORT_PROMPT + "\n\nDocument Content for Analysis:\n"]
        for rel_item, meta in relevant:
            item = rel_item
            file_name, _ = meta
            if item['type'] == 'text':
                content_list.append(f"Relevant Text from {file_name} (Page {item['metadata'].get('page', 0)}, Chunk {item['metadata'].get('chunk_idx', 0)}):\n{item['content']}\n")
            elif item['type'] == 'image':
                position = item['metadata']['position']
                content_list.append(f"Relevant Image from {file_name} (Page {item['metadata']['page']}, Position {position}):\n")
                content_list.append({
                    'inline_data': {
                        'mime_type': 'image/png',
                        'data': base64.b64encode(item['content']).decode()
                    }
                })
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(content_list)
            report_text = response.text
            st.markdown("### Generated Analysis Report")
            st.markdown(report_text)
            word_filename = save_report_to_word(report_text)
            if word_filename:
                with open(word_filename, "rb") as file:
                    st.download_button(
                        label="Download Generated Analysis Report as Word Document",
                        data=file,
                        file_name=word_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")
else:
    if not uploaded_files:
        st.info("Please upload at least one file to generate the report.")

# Chat Input for Q&A
st.markdown("### Ask RE Analyst")
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        if message["role"] == "assistant" and os.path.exists(LOGO_PATH):
            st.image(LOGO_PATH, width=50)
        st.markdown(message["content"])

user_query = st.chat_input("Ask a question about the files:")
if user_query and faiss_index:
    with st.chat_message("user"):
        st.markdown(user_query)
    st.session_state.chat_history.append({"role": "user", "content": user_query})

    with st.spinner("Generating answer..."):
        relevant = retrieve_relevant_chunks(user_query, faiss_index, items_list, chunk_metadata)
        content_list = [QNA_PROMPT + "\n\nRelevant Document Content:\n"]
        for rel_item, meta in relevant:
            item = rel_item
            file_name, _ = meta
            if item['type'] == 'text':
                content_list.append(f"Relevant Text from {file_name} (Page {item['metadata'].get('page', 0)}, Chunk {item['metadata'].get('chunk_idx', 0)}):\n{item['content']}\n")
            elif item['type'] == 'image':
                position = item['metadata']['position']
                content_list.append(f"Relevant Image from {file_name} (Page {item['metadata']['page']}, Position {position}):\n")
                content_list.append({
                    'inline_data': {
                        'mime_type': 'image/png',
                        'data': base64.b64encode(item['content']).decode()
                    }
                })
        content_list.append(f"\nUser Question: {user_query}")
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(content_list)
            answer_text = response.text
            with st.chat_message("assistant"):
                if os.path.exists(LOGO_PATH):
                    st.image(LOGO_PATH, width=50)
                st.markdown(answer_text)
            st.session_state.chat_history.append({"role": "assistant", "content": answer_text})
        except Exception as e:
            st.error(f"Error generating answer: {str(e)}")
else:
    if not uploaded_files:
        st.info("Please upload at least one file to begin Q&A.")
    if not faiss_index and uploaded_files:
        st.info("Document indexing failed. Check error messages above.")

# Footer
st.sidebar.markdown("---")
st.sidebar.info("This app uses Google Gemini 1.5 Flash, CLIP, and FAISS for multimodal document analysis. For support, check the Setup Instructions above.")
